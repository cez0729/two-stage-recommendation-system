"""Create the final Chinese research report from verified project artifacts."""

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "final_research_report_zh.docx"


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values, strict=True):
            cell.text = text
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(value)


def format_mean_sd(summary: dict, metric: str, *, percent: bool = True) -> str:
    mean = summary[metric]["mean"]
    std = summary[metric]["std_sample"]
    if percent:
        return f"{mean:.2%} ± {std:.2%}"
    return f"{mean:.2f} ± {std:.2f}"


def main() -> None:
    two_tower = json.loads(
        (ROOT / "results/video_games_2018/two_tower_v3_final_metrics.json").read_text(
            encoding="utf-8"
        )
    )
    final_test = json.loads(
        (ROOT / "results/video_games_2018/final_test_metrics.json").read_text(encoding="utf-8")
    )
    serving_benchmark = json.loads(
        (ROOT / "results/video_games_2018/serving_benchmark_profiled.json").read_text(
            encoding="utf-8"
        )
    )
    stage_profile = json.loads(
        (ROOT / "results/video_games_2018/serving_stage_profile.json").read_text(
            encoding="utf-8"
        )
    )
    phase6_dir = ROOT / "results/video_games_2018/phase6"
    metadata_audit = json.loads((phase6_dir / "metadata_audit.json").read_text(encoding="utf-8"))
    content_experiment = json.loads(
        (phase6_dir / "content_experiment.json").read_text(encoding="utf-8")
    )
    global_temporal = json.loads(
        (phase6_dir / "global_temporal_experiment.json").read_text(encoding="utf-8")
    )
    logq_ablation = json.loads(
        (phase6_dir / "logq_ablation_summary.json").read_text(encoding="utf-8")
    )
    load_test = json.loads((phase6_dir / "load_test.json").read_text(encoding="utf-8"))
    feature_build_share = stage_profile["stages"]["feature_build"][
        "share_of_total_mean_pct"
    ]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("两阶段推荐系统\n最终研究报告")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph("Amazon Video Games · 离线研究原型与可观测 Serving Demo")
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    add_table(
        doc,
        ["项目状态", "数据规模", "评估口径", "报告日期"],
        [["Phase 1–6 完成", "197,597 交互 / 19,621 用户", "完整目录 + 时间稳健性", "2026-09-01"]],
        [2100, 2500, 2500, 2260],
    )
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(
        "本项目已经完成一个可复现的两阶段推荐系统研究原型，并进一步封装为可部署、"
        "可监控的本地 Serving Demo。系统先用 Popularity、ItemCF 和修复后的 Two-Tower "
        "生成候选，再用 FAISS、RRF 和 LightGBM LambdaRank 完成排序。Phase 6 增加 title、brand、"
        "细粒度类目内容召回，并完成冷启动、全局 cutoff、三随机种子和并发实验。"
        "冻结的三通道 baseline 使用 16 个排序特征；content-aware ranker 新增 content_score "
        "和 content_rank，"
        "共 18 个特征。"
        "所有最终数字来自真实运行产物；"
        "项目没有使用线上曝光日志，因此不宣称真实商业收益。"
    )
    doc.add_heading("关键结论", level=2)
    add_bullets(
        doc,
        [
            "Two-Tower 的测试 Recall@100 为 9.51%，超过 Popularity 的 6.41%，"
            "但仍低于 ItemCF 的 15.70%。",
            "多路候选测试 Recall@200 为 21.00%；LambdaRank 将测试 NDCG@10 从 RRF 的 1.78% "
            "提升到 2.99%。",
            "在统一的内容消融协议下，候选预算仍为 200；加入内容通道后测试 "
            "Candidate Recall@200 从 matched 三通道基线 20.65% "
            "提升到 24.20%；721 位严格冷用户的 Content Recall@100 为 14.29%，Two-Tower 为 0%。",
            "三随机种子都支持 logQ 的准确率收益，但 Coverage@10 从约 99.10% 降到 68.31%，"
            "热门偏置比从 0.40 升到 2.09，必须作为权衡披露。",
            "FAISS IndexFlatIP 在 100 位用户上的 Top-100 与 NumPy 精确点积 100% 一致，"
            "单用户 P95 为 2.15 ms。",
            f"顺序 Serving 微基准 P50/P95 为 {serving_benchmark['p50_ms']:.2f}/"
            f"{serving_benchmark['p95_ms']:.2f} ms，但并发压测显示单进程仅约 11–15 QPS，"
            "不能把顺序请求速率当作容量。分阶段测量显示特征构造是最大耗时项，"
            f"占平均总耗时 {feature_build_share:.2f}%。",
        ],
    )
    doc.add_heading("1. 数据与评估设计", level=1)
    doc.add_paragraph(
        "正式数据来自 UCSD McAuley Lab 的 Amazon Reviews 2018 Video_Games 5-core。"
        "原始评论按用户、商品和日期整理，只保留正向隐式反馈，并为每位用户按不同日期构造 "
        "retrieval_train、rank_train、validation 和 test。主召回指标使用完整 13,923 商品训练目录，"
        "而不是 sampled-negative 评估。"
    )
    add_table(
        doc,
        ["对象", "数量", "用途"],
        [
            ["原始评论", "497,577", "数据下载与审计"],
            ["正式交互", "197,597", "清洗、切分与训练"],
            ["可评估用户", "19,621", "每人一个 validation/test 目标"],
            ["商品总数", "14,391", "清洗后商品实体"],
            ["训练召回目录", "13,923", "双塔、ItemCF、FAISS 共同目录"],
        ],
        [2400, 1800, 5160],
    )
    doc.add_heading("2. 失败诊断与模型修复", level=1)
    doc.add_paragraph(
        "第一版双塔在测试集只有 Recall@100=3.02%，低于 Popularity。代码检查证明 checkpoint、"
        "向量归一化、全目录检索和已见过滤均正常。继续训练到 30 轮后验证集也只有 4.88%，"
        "因此问题不是简单欠训练。"
    )
    doc.add_paragraph(
        "进一步分析发现，批内负样本按出现频率被抽取，热门商品更容易成为负例，"
        "未校正模型推荐商品平均热度只有真实目标的 39%。加入 logQ 采样校正后，双塔验证 Recall@100 "
        "达到 11.48%，冻结测试为 9.51%。这项修复在未参与调参的测试集上仍有效。"
    )
    doc.add_heading("3. 最终离线结果", level=1)
    tt_test = two_tower["splits"]["test"]["metrics"]
    add_table(
        doc,
        ["模型/阶段", "Recall@10", "Recall@100", "NDCG@10", "说明"],
        [
            ["Popularity", "—", "6.41%", "—", "最简单基线"],
            ["ItemCF", "—", "15.70%", "—", "最强单路召回基线"],
            [
                "Two-Tower + logQ",
                f"{tt_test['recall@10']:.2%}",
                f"{tt_test['recall@100']:.2%}",
                f"{tt_test['ndcg@10']:.2%}",
                "完整目录测试",
            ],
            [
                "多路 RRF",
                f"{final_test['retrieval_order_metrics']['recall@10']:.2%}",
                "21.00%候选召回",
                f"{final_test['retrieval_order_metrics']['ndcg@10']:.2%}",
                "200 候选",
            ],
            [
                "多路 + LambdaRank",
                f"{final_test['lambdarank_metrics']['recall@10']:.2%}",
                "21.00%候选召回",
                f"{final_test['lambdarank_metrics']['ndcg@10']:.2%}",
                "冻结测试",
            ],
        ],
        [2600, 1300, 1500, 1300, 2660],
    )
    doc.add_paragraph(
        "LambdaRank 的测试 Recall@10 绝对提升 2.06 个百分点，95% 成对 bootstrap 区间为 "
        "[1.77, 2.34]；NDCG@10 绝对提升 1.21 个百分点，区间为 [1.04, 1.37]。新增命中 623 人、"
        "损失 219 人，McNemar 精确检验 p≈1.3×10^-45。"
    )
    doc.add_heading("4. 工程化 Serving", level=1)
    doc.add_paragraph(
        "ServingPipeline 在启动时一次性加载冻结模型、FAISS、ItemCF、Popularity、"
        "Content 和 LambdaRank。"
        "known user 使用完整流程；unknown user 使用稳定的 Popularity fallback。API 只暴露 /health、"
        "/recommend/{user_id} 和 /metrics，k 限制为 1–50。Redis 使用 "
        "rec:v1:{user_id}:{k}、TTL 300 秒，"
        "连接失败在 50 ms 后自动绕过。"
    )
    add_table(
        doc,
        ["验收项", "结果"],
        [
            ["离线 parity", "20/20 固定用户 top-10 完全一致"],
            ["FAISS 正确性", "100 位用户逐行 100% 匹配 NumPy"],
            ["重复请求", "结果确定性一致；无重复、无已见商品"],
            [
                "Pipeline profiling",
                f"P95 {stage_profile['total_pipeline']['p95_ms']:.2f} ms；"
                f"候选数 {stage_profile['candidate_count']['min']}–"
                f"{stage_profile['candidate_count']['max']}",
            ],
            [
                "API 顺序微基准",
                f"P50 {serving_benchmark['p50_ms']:.2f} ms / "
                f"P95 {serving_benchmark['p95_ms']:.2f} ms / "
                f"{serving_benchmark['qps']:.2f} req/s（非容量）",
            ],
            ["Docker", "Phase 6 已在 AWS 8001 独立测试栈验收，8000 基线保持健康"],
        ],
        [2600, 6760],
    )
    doc.add_heading("5. Phase 6 元数据门禁", level=1)
    audit_catalog = metadata_audit["fields"]
    audit_cold = metadata_audit["strict_cold"]["coverage"]
    add_table(
        doc,
        ["字段", "全目录覆盖率", "严格冷目标覆盖率", "决策"],
        [
            [
                "Title",
                f"{audit_catalog['title']['coverage']:.2%}",
                f"{audit_cold['title']:.2%}",
                "使用",
            ],
            [
                "Brand",
                f"{audit_catalog['brand']['coverage']:.2%}",
                f"{audit_cold['brand']:.2%}",
                "使用",
            ],
            [
                "细粒度类目",
                f"{audit_catalog['fine_category']['coverage']:.2%}",
                f"{audit_cold['fine_category']:.2%}",
                "使用",
            ],
            ["Price/聚合评分", "不可靠", "不可靠", "排除"],
        ],
        [2100, 2100, 2300, 2860],
    )
    doc.add_paragraph(
        f"元数据质量门禁结论为 {metadata_audit['decision']}。TF-IDF 只使用 title、规范化 brand 和"
        "细粒度类目路径；词表只在 retrieval-train 可见商品上拟合，"
        "但检索目录包含所有元数据可见商品。"
        "当前混合是 score-level RRF，不是训练过的 Hybrid Item Tower。"
    )
    doc.add_heading("6. 内容召回与严格冷启动", level=1)
    validation_metrics = content_experiment["splits"]["validation"]["metrics"]
    test_metrics = content_experiment["splits"]["test"]["metrics"]
    add_table(
        doc,
        ["指标", "验证", "测试", "解释"],
        [
            [
                "Content Recall@100",
                f"{validation_metrics['content']['recall@100']:.2%}",
                f"{test_metrics['content']['recall@100']:.2%}",
                "只依赖商品文本元数据",
            ],
            [
                "Two-Tower Recall@100",
                f"{validation_metrics['two_tower']['recall@100']:.2%}",
                f"{test_metrics['two_tower']['recall@100']:.2%}",
                "协同神经召回",
            ],
            [
                "TT+Content Recall@100",
                f"{validation_metrics['hybrid_rrf_two_tower_content']['recall@100']:.2%}",
                f"{test_metrics['hybrid_rrf_two_tower_content']['recall@100']:.2%}",
                "分数级 RRF",
            ],
            [
                "内容消融 matched 三通道 CR@200",
                f"{validation_metrics['three_collaborative']['candidate_recall@200']:.2%}",
                f"{test_metrics['three_collaborative']['candidate_recall@200']:.2%}",
                "固定 200 候选",
            ],
            [
                "四通道 CR@200",
                f"{validation_metrics['all_four_channels']['candidate_recall@200']:.2%}",
                f"{test_metrics['all_four_channels']['candidate_recall@200']:.2%}",
                "固定 200 候选",
            ],
        ],
        [2600, 1400, 1400, 3960],
    )
    cold_comparison = content_experiment["splits"]["test"]["primary_comparisons"][
        "content_minus_two_tower_strict_cold_r100"
    ]
    cold_ci = cold_comparison["left_minus_right"]
    cold_mcnemar = cold_comparison["mcnemar"]
    doc.add_paragraph(
        f"测试严格冷用户共 {cold_comparison['users']} 位，Content 命中 "
        f"{cold_mcnemar['left_only_hits']} 位，Recall@100={cold_comparison['left_recall']:.2%}；"
        f"Two-Tower={cold_comparison['right_recall']:.2%}。差值 95% CI "
        f"[{cold_ci['ci95_low']:.2%}, {cold_ci['ci95_high']:.2%}]，McNemar 精确检验 "
        f"p={cold_mcnemar['exact_p_value']:.2e}。四通道在测试中新增 1,079 个命中、损失 382 个命中，"
        "净增不是通过扩大候选预算获得。"
    )
    doc.add_heading("7. 全局时间稳健性", level=1)
    global_final = global_temporal["periods"]["final_holdout"]
    global_metrics = global_final["metrics"]
    global_cold = global_final["comparisons"]["content_minus_itemcf_strict_cold_r100"]
    add_table(
        doc,
        ["2017 cutoff 最终留出", "Content/融合", "协同基线", "差值含义"],
        [
            [
                "Overall Recall@100",
                f"{global_metrics['content']['recall@100']:.2%}",
                f"{global_metrics['itemcf']['recall@100']:.2%}",
                "Content 对 ItemCF",
            ],
            [
                "严格冷 Recall@100",
                f"{global_cold['left_recall']:.2%}",
                f"{global_cold['right_recall']:.2%}",
                f"{global_cold['users']} 位冷用户",
            ],
            [
                "Candidate Recall@200",
                f"{global_metrics['hybrid_content_collaborative']['candidate_recall@200']:.2%}",
                f"{global_metrics['two_collaborative']['candidate_recall@200']:.2%}",
                "Content+协同 对 仅协同",
            ],
        ],
        [2500, 1800, 1800, 3260],
    )
    doc.add_paragraph(
        f"验证 cutoff={global_temporal['protocol']['validation_cutoff']}，最终 cutoff="
        f"{global_temporal['protocol']['final_cutoff']}，日期在查看指标前固定。"
        "旧 Two-Tower checkpoint "
        "因见过更晚数据而被排除。结论方向在统一时间切分下仍成立；但是静态元数据的字段级可用时间"
        "无法独立验证，因此这是一项稳健性检查，不是完全无偏的历史回放。"
    )
    doc.add_heading("8. 三随机种子 logQ 消融", level=1)
    raw_summary = logq_ablation["summary"]["raw"]
    logq_summary = logq_ablation["summary"]["logq"]
    add_table(
        doc,
        ["指标", "Raw 均值 ± SD", "logQ 均值 ± SD", "三种子方向"],
        [
            [
                "Recall@100",
                format_mean_sd(raw_summary, "recall@100"),
                format_mean_sd(logq_summary, "recall@100"),
                "全部提升",
            ],
            [
                "Recall@10",
                format_mean_sd(raw_summary, "recall@10"),
                format_mean_sd(logq_summary, "recall@10"),
                "全部提升",
            ],
            [
                "NDCG@10",
                format_mean_sd(raw_summary, "ndcg@10"),
                format_mean_sd(logq_summary, "ndcg@10"),
                "全部提升",
            ],
            [
                "Coverage@10",
                format_mean_sd(raw_summary, "coverage@10"),
                format_mean_sd(logq_summary, "coverage@10"),
                "全部下降",
            ],
            [
                "热门偏置比",
                format_mean_sd(raw_summary, "popularity_bias_ratio", percent=False),
                format_mean_sd(logq_summary, "popularity_bias_ratio", percent=False),
                "全部上升",
            ],
        ],
        [2100, 2500, 2500, 2260],
    )
    doc.add_paragraph(
        "受控实验表明 logQ 的准确率修复可重复，但不是所有维度全面胜出。更高准确率伴随更低目录覆盖和"
        "更强热门偏置；未来应在验证集上联合优化 Recall、Coverage、Novelty 和 Diversity。"
    )
    doc.add_heading("9. 并发压测与容量边界", level=1)
    load_rows = []
    for level in load_test["levels"]:
        load_rows.append(
            [
                str(level["concurrency"]),
                f"{level['qps']:.2f}",
                f"{level['p95_ms']:.2f} ms",
                f"{level['p99_ms']:.2f} ms",
                str(level["errors"]),
            ]
        )
    add_table(
        doc,
        ["并发", "QPS", "P95", "P99", "错误"],
        load_rows,
        [1300, 1600, 2300, 2300, 1860],
    )
    doc.add_paragraph(
        "环境为单个本地 Uvicorn 进程、无 Redis、100 位已知用户闭环请求。吞吐在约 11–15 QPS 饱和，"
        "并发增加时尾延迟迅速放大。零错误说明功能路径稳定，但不能据此宣称生产容量或 SLA。"
    )
    load_plot = phase6_dir / "load_test.png"
    if load_plot.exists():
        doc.add_picture(str(load_plot), width=Inches(6.2))
        caption = doc.add_paragraph("图 1  单进程并发压测：吞吐平台与尾延迟增长")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
    doc.add_heading("10. 限制与下一步", level=1)
    add_bullets(
        doc,
        [
            "数据是公开 Amazon Reviews，只有离线反馈，没有曝光日志、线上反馈或 A/B 实验。",
            "内容召回已改善严格冷商品，但 TF-IDF 和分数级 RRF 仍不是联合训练的 Hybrid Item Tower。",
            "静态元数据的历史可用时间无法独立验证；price、avg_rating、rating_number 继续排除。",
            "单进程并发扩展性差；Phase 6 已完成 20/20 parity，"
            "并在 AWS 8001 独立测试栈通过容器验收。",
            "下一步应在验证集学习融合权重并重训含内容特征的 ranker，"
            "同时加入 novelty/diversity 约束。",
        ],
    )
    heading = doc.add_heading("11. 项目定位与简历表述", level=1)
    heading.paragraph_format.keep_with_next = True
    doc.add_paragraph(
        "准确定位：End-to-End Production-Style Two-Stage Recommendation System，"
        "可部署、可监控的离线研究原型。它适合用于研究生申请和推荐算法实习，含金量来自失败诊断、"
        "防泄漏评估、采样校正、ANN 正确性、排序显著性和可观测 Serving，而不是虚构线上业务收益。"
    )
    doc.add_paragraph(
        "推荐简历表述：Built a leakage-aware two-stage recommender on 197K temporally split "
        "interactions; diagnosed in-batch sampling bias with a three-seed logQ ablation, verified "
        "exact FAISS retrieval, and added metadata-based content retrieval that improved "
        "fixed-budget Candidate Recall@200 from 20.65% to 24.20% while reaching 14.29% "
        "Recall@100 on 721 strict-cold users."
    )
    doc.add_paragraph(
        "面试时应把证据分成三层：第一层是离线研究证据（完整目录 Recall、冷启动差值、"
        "bootstrap/McNemar、三种子方向）；第二层是工程证据（模型与数据哈希、FAISS exact match、"
        "20/20 parity、并发曲线）；第三层是尚未具备的业务证据（曝光、点击、购买、A/B、"
        "CTR/CVR/GMV）。"
        "前两层可以明确陈述，第三层必须说明尚不存在。"
    )
    doc.add_paragraph(
        "30 秒中文讲法：我在 19.8 万条公开 Amazon 交互上做了防泄露两阶段推荐，"
        "先用多路召回把全目录缩到 200 个候选，再用 LambdaRank 排序。双塔异常低分后，"
        "我定位并修复了批内负采样偏差；针对协同模型无法处理的新商品，我通过元数据门禁后"
        "加入 TF-IDF 内容召回，在固定候选预算下把测试 Candidate Recall@200 "
        "从 20.65% 提升到 24.20%，"
        "严格冷用户 Recall@100 达到 14.29%。我还做了全局时间稳健性、三随机种子"
        "消融和并发压测，并明确没有把离线结果包装成线上商业收益。"
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Two-stage Recommendation System | Final Research Report")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
